from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "MedTrace_AI_Technical_Summary.docx"

TEAL = "168985"
NAVY = "071B25"
MUTED = "64767C"
LIGHT_TEAL = "E8F7F5"
LIGHT_GRAY = "F4F7F7"
LINE = "DCE6E7"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color: str = LINE) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = "w:{}".format(edge)
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_text(cell, text: str, bold: bool = False, color: str | None = None) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(9.2)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_heading(doc: Document, text: str, level: int = 1):
    p = doc.add_paragraph()
    p.style = "Heading {}".format(level)
    run = p.add_run(text)
    run.font.color.rgb = RGBColor.from_string(NAVY if level == 1 else TEAL)
    return p


def add_body(doc: Document, text: str):
    p = doc.add_paragraph()
    p.style = "Body Text"
    p.add_run(text)
    return p


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)


def add_callout(doc: Document, title: str, body: str, fill: str = LIGHT_TEAL) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_border(cell, "B9DAD6")
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(title)
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(TEAL)
    r.font.size = Pt(10.5)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    r2 = p2.add_run(body)
    r2.font.size = Pt(9.3)
    r2.font.color.rgb = RGBColor.from_string(NAVY)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for i, header in enumerate(headers):
        cell = hdr.cells[i]
        set_cell_shading(cell, TEAL)
        set_cell_border(cell, "0F6F6B")
        set_cell_text(cell, header, bold=True, color="FFFFFF")
        if widths:
            cell.width = Cm(widths[i])
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_border(cells[i])
            set_cell_text(cells[i], value)
            if widths:
                cells[i].width = Cm(widths[i])
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(1.7)
    section.bottom_margin = Cm(1.65)
    section.left_margin = Cm(1.75)
    section.right_margin = Cm(1.75)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos")
    normal.font.size = Pt(10)

    body = styles["Body Text"]
    body.font.name = "Aptos"
    body._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos")
    body.font.size = Pt(10)
    body.paragraph_format.line_spacing = 1.08
    body.paragraph_format.space_after = Pt(7)

    for i, size in [(1, 18), (2, 14), (3, 11.5)]:
        style = styles["Heading {}".format(i)]
        style.font.name = "Aptos Display"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos Display")
        style.font.bold = True
        style.font.size = Pt(size)
        style.paragraph_format.space_before = Pt(11 if i == 1 else 8)
        style.paragraph_format.space_after = Pt(5)

    bullet = styles["List Bullet"]
    bullet.font.name = "Aptos"
    bullet.font.size = Pt(9.6)
    bullet.paragraph_format.space_after = Pt(3)


def build() -> None:
    doc = Document()
    configure_document(doc)
    props = doc.core_properties
    props.title = "MedTrace AI Technical Summary"
    props.subject = "YGC AI Competition 2026 technical summary"
    props.author = "MedTrace AI Team"
    props.comments = "Editable technical summary for submission preparation."
    props.keywords = "MedTrace AI, medical records, Supabase, n8n, OpenAI, evidence citations"

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("MedTrace AI")
    r.bold = True
    r.font.name = "Aptos Display"
    r.font.size = Pt(28)
    r.font.color.rgb = RGBColor.from_string(NAVY)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = subtitle.add_run("Technical Summary for YGC AI Competition 2026")
    r.font.size = Pt(13)
    r.font.color.rgb = RGBColor.from_string(TEAL)

    tagline = doc.add_paragraph()
    tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = tagline.add_run("Every record in view. Every claim anchored.")
    r.italic = True
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor.from_string(MUTED)

    add_callout(
        doc,
        "Executive summary",
        "MedTrace AI is an evidence-linked medical record intelligence platform. It organizes fragmented prescriptions, lab reports, and clinical notes into one patient timeline, highlights possible safety candidates, and answers record-level questions with document/page/snippet citations. The public demo is stable and credential-free; the connected architecture is prepared for Supabase, n8n, and OpenAI-based processing.",
    )

    add_heading(doc, "1. Problem And Product Goal", 1)
    add_body(
        doc,
        "Patients often collect medical records across different visits, providers, and formats. Important details such as allergies, active medications, dosage changes, lab values, and clinical notes can become separated. MedTrace AI addresses this fragmentation by creating an evidence-linked workspace where every important claim remains connected to its source page.",
    )
    add_body(
        doc,
        "The goal is not to diagnose or prescribe. The platform supports careful review by showing what the uploaded record appears to contain, which claims need professional attention, and exactly where each claim came from.",
    )
    add_bullets(
        doc,
        [
            "Create and open patient workspaces.",
            "Review a chronological timeline across many documents.",
            "Inspect medications, allergies, laboratory values, diagnoses, visits, and providers.",
            "Detect possible duplicate prescriptions, dosage conflicts, allergy contradictions, and curated interaction candidates.",
            "Ask questions across multiple documents and receive cited, record-grounded answers.",
            "Separate clinical risk severity from evidence confidence.",
        ],
    )

    add_heading(doc, "2. Technical Stack", 1)
    add_table(
        doc,
        ["Layer", "Technology", "Purpose"],
        [
            ["Frontend", "React, TypeScript, Vite, Tailwind CSS", "Responsive web application, patient dashboard, evidence viewer, demo flow."],
            ["State and data", "TanStack Query, adapter boundary", "Keeps demo mode and Supabase-connected mode behind one client interface."],
            ["Charts", "Recharts", "Laboratory trend visualization."],
            ["Auth and database", "Supabase Auth, PostgreSQL, RLS", "Authenticated users, patient-owned rows, relational clinical data."],
            ["Storage and retrieval", "Private Supabase Storage, pgvector", "Private document storage and patient-filtered semantic retrieval."],
            ["Workflow backend", "n8n webhooks and sub-workflows", "Visible multi-agent orchestration, retries, audit records, and API endpoints."],
            ["AI", "OpenAI Responses API and embeddings", "Document classification, structured extraction, evidence verification, explanations, Q&A."],
            ["Rules", "TypeScript deterministic clinical rules", "Duplicate, dosage, allergy, confidence, citation, ownership, and safety checks."],
            ["Testing", "Vitest, Playwright, TypeScript build", "Unit tests, browser journey tests, type safety, production build verification."],
        ],
        [2.7, 4.2, 8.4],
    )

    add_heading(doc, "3. System Architecture", 1)
    add_body(
        doc,
        "The system intentionally avoids a separate Express or FastAPI server. The browser communicates directly with Supabase for authenticated reads and private uploads, while long-running processing and record Q&A are handled by authenticated n8n production webhooks.",
    )
    add_table(
        doc,
        ["Step", "Runtime Flow"],
        [
            ["1", "The user signs in or opens the controlled demo path."],
            ["2", "The browser creates or opens a patient workspace."],
            ["3", "In connected mode, documents upload to private Supabase Storage using the user's session."],
            ["4", "The browser calls n8n webhooks for processing or Q&A with patient and job identifiers."],
            ["5", "n8n validates JWT signature, expiry, audience, and patient ownership before work begins."],
            ["6", "n8n invokes bounded sub-workflows for classification, extraction, normalization, reconciliation, trend analysis, safety rules, evidence verification, and Q&A."],
            ["7", "Server-side Supabase credentials write normalized rows, audit records, findings, chunks, and failures."],
            ["8", "The frontend reads patient-owned results through RLS and displays only cited claims."],
        ],
        [1.3, 14.0],
    )
    add_callout(
        doc,
        "Trust boundary",
        "The browser may receive only the Supabase anon key and the signed-in user's JWT. The service-role key, OpenAI API key, n8n credentials, and webhook secret must remain server-side.",
        LIGHT_GRAY,
    )

    add_heading(doc, "4. AI And Multi-Agent Workflow Design", 1)
    add_body(
        doc,
        "The AI pipeline is designed as a set of bounded n8n workflows rather than one unrestricted autonomous agent. Each workflow has one responsibility, typed JSON inputs and outputs, versioned prompts, timeout/retry limits, and an audit trail.",
    )
    add_table(
        doc,
        ["Workflow", "Responsibility"],
        [
            ["WF-01 Process Patient Record", "Public intake webhook, contract validation, ownership checks, job orchestration."],
            ["WF-02 Process One Document", "Page-preserving document classification and clinical extraction."],
            ["WF-03 Normalize Medication", "Medication name normalization through local aliases/RxNorm adapter."],
            ["WF-04 Normalize Laboratory Result", "Unit-aware lab parsing and comparison eligibility."],
            ["WF-05 Medication Reconciliation", "Duplicate and dosage-conflict candidate preparation."],
            ["WF-06 Laboratory Trends", "Compatible-unit trend grouping and explanation."],
            ["WF-07 Safety Rules", "Deterministic duplicate, allergy, dosage, and cited interaction candidates."],
            ["WF-08 Evidence Verification", "Checks that each displayed finding is supported by source evidence."],
            ["WF-09 Ask Record Question", "Patient-scoped retrieval and cited Q&A."],
            ["WF-99 Error Handler", "Sanitized workflow failure capture and retry policy."],
        ],
        [4.5, 10.8],
    )
    add_body(
        doc,
        "OpenAI is used for multimodal understanding, structured extraction, evidence verification, restrained explanations, and record-grounded Q&A. Deterministic rules remain responsible for safety candidate generation where precision and auditability matter.",
    )

    add_heading(doc, "5. Evidence, Safety, And Clinical Boundary", 1)
    add_body(
        doc,
        "Every extracted entity and derived finding retains provenance: document ID, page number, exact supporting snippet, confidence, prompt version, and model metadata. The application displays evidence confidence separately from clinical risk severity.",
    )
    add_table(
        doc,
        ["Safety Principle", "Implementation"],
        [
            ["No diagnosis or prescribing", "The UI states that the app does not diagnose and must not be used to start, stop, or change medication."],
            ["Evidence-first display", "Findings and answers are displayed only with document/page/snippet citations."],
            ["Risk is separate from confidence", "Risk indicates possible importance; confidence indicates evidence quality."],
            ["Prompt-injection resistance", "Uploaded document text is treated as untrusted input and cannot override system behavior."],
            ["Patient isolation", "RLS, patient ownership checks, and citation validation prevent cross-patient leakage."],
            ["Interaction truth boundary", "Drug interaction truth must come from deterministic curated/licensed providers, not solely from a language model."],
        ],
        [4.0, 11.2],
    )

    add_heading(doc, "6. Database And Supabase Design", 1)
    add_body(
        doc,
        "Supabase provides authentication, relational storage, private file storage, Row Level Security, Realtime job updates, and pgvector retrieval. The migration set creates the core clinical schema, indexes, storage policy, search RPC, and cross-patient integrity triggers.",
    )
    add_bullets(
        doc,
        [
            "Core tables include patients, documents, document_pages, document_chunks, visits, medications, allergies, lab_results, diagnoses, timeline_events, findings, finding_evidence, processing_jobs, qa_threads, qa_messages, agent_runs, and workflow_failures.",
            "RLS policies ensure users can access only their own patient workspaces.",
            "A private medical-documents bucket prevents permanent public file URLs.",
            "A patient-filtered vector-search RPC supports grounded retrieval without crossing ownership boundaries.",
            "Idempotency keys and file hashes reduce duplicate processing.",
        ],
    )

    add_heading(doc, "7. Current Implementation Status", 1)
    add_table(
        doc,
        ["Area", "Status"],
        [
            ["Frontend", "Implemented and polished: landing page, patient list, dashboard, tabs, findings, evidence drawer, Q&A, upload/reprocess dialogs, mobile responsiveness."],
            ["Demo path", "Implemented: cached synthetic competition walkthrough works without Supabase, n8n, or OpenAI credentials."],
            ["Supabase", "SQL migrations are ready and the project owner has created the Supabase project/tables."],
            ["n8n", "Ten workflow templates are exported and ready to import/configure."],
            ["AI prompts", "Versioned prompts exist for classifier, extractor, reconciler, lab trend, verifier, and Q&A."],
            ["Tests", "Lint, typecheck, unit tests, and production build pass."],
            ["Deployment", "Frontend has been deployed to Vercel; final environment variables and n8n import remain."],
        ],
        [3.5, 11.8],
    )

    add_heading(doc, "8. Evaluation Strategy", 1)
    add_body(
        doc,
        "Evaluation focuses on citation validity, schema validity, medication precision, lab numeric accuracy, safety-candidate recall on reviewed cases, and refusal of unsupported answers. The project also includes adversarial checks for duplicate files, prompt-like document text, missing citations, cross-patient citation attempts, incompatible lab units, and ambiguous medication/allergy aliases.",
    )
    add_table(
        doc,
        ["Metric", "Target"],
        [
            ["Structured schema validity", "100%"],
            ["Citation validity", "100%"],
            ["Unsupported critical findings", "0"],
            ["Safety-candidate recall on reviewed cases", "100%"],
            ["Medication precision", "At least 90%"],
            ["Laboratory numeric accuracy", "At least 95%"],
            ["Grounded Q&A answers", "At least 90%"],
        ],
        [6.5, 8.8],
    )

    add_heading(doc, "9. Challenges Faced", 1)
    add_bullets(
        doc,
        [
            "Medical records can be incomplete, duplicated, inconsistent, or mislabeled; the system therefore treats filenames and folder labels as untrusted.",
            "A useful demo must be reliable even when external services are not ready; the application keeps a credential-free cached demo mode.",
            "Clinical safety requires controlled reasoning, so deterministic rules and evidence verification are used before displaying findings.",
            "Supabase and n8n credentials must be separated carefully so service-role keys and OpenAI keys never reach the browser.",
            "Mobile dashboards can easily become too wide; the frontend was tightened with responsive tabs, stacked actions, and safer evidence views.",
        ],
    )

    add_heading(doc, "10. Unique Impact And Value Proposition", 1)
    add_body(
        doc,
        "MedTrace AI is valuable because it does not simply summarize documents. It creates a traceable evidence layer. A reviewer can see a possible issue, open the exact supporting source page, understand the confidence level, and ask follow-up questions without losing the connection to the underlying record.",
    )
    add_callout(
        doc,
        "Core value",
        "The platform makes fragmented medical records easier to navigate while preserving the humility required in clinical contexts: every claim is anchored, unsupported questions are refused, and high-risk or uncertain findings recommend professional verification.",
    )

    add_heading(doc, "11. Final Submission Readiness", 1)
    add_bullets(
        doc,
        [
            "Public frontend: deployed on Vercel.",
            "Database: Supabase project created and SQL tables prepared.",
            "Demo: stable cached walkthrough ready for judging.",
            "Pending setup: configure production environment variables and import/configure the ten n8n workflows.",
            "Presentation path: open demo, inspect patient dashboard, review a high-risk finding, open evidence, ask a cited question, and show insufficient-evidence behavior.",
        ],
    )

    section = doc.add_section(WD_SECTION.CONTINUOUS)
    section.top_margin = Cm(1.7)
    section.bottom_margin = Cm(1.65)
    section.left_margin = Cm(1.75)
    section.right_margin = Cm(1.75)

    add_heading(doc, "Appendix: Environment Checklist", 1)
    add_table(
        doc,
        ["Variable", "Where It Belongs", "Purpose"],
        [
            ["VITE_DATA_MODE", "Frontend", "demo for reliable public demo; supabase after full integration."],
            ["VITE_SUPABASE_URL", "Frontend", "Public Supabase project URL."],
            ["VITE_SUPABASE_ANON_KEY", "Frontend", "Supabase anon key protected by RLS."],
            ["VITE_N8N_BASE_URL", "Frontend", "Public n8n base URL for connected mode."],
            ["SUPABASE_SERVICE_ROLE_KEY", "Server/n8n only", "Privileged database writes from workflow backend."],
            ["SUPABASE_JWT_SECRET", "Server/n8n only", "JWT verification secret for legacy HS256 validation."],
            ["OPENAI_API_KEY", "Server/n8n only", "AI extraction, verification, embeddings, and Q&A."],
            ["N8N_WEBHOOK_SECRET", "Server/n8n only", "Shared webhook protection secret."],
        ],
        [4.4, 4.0, 6.9],
    )

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = footer.add_run("MedTrace AI - Technical Summary - YGC AI Competition 2026")
    fr.font.size = Pt(8)
    fr.font.color.rgb = RGBColor.from_string(MUTED)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
